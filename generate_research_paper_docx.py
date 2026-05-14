from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


OUTPUT_PATH = "AWANSHI_Enterprises_Research_Paper.docx"


def set_page_number(paragraph):
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    heading.style.font.name = "Times New Roman"
    heading.style.font.size = Pt(14 if level == 1 else 12)
    return heading


def add_paragraph(doc, text, bold=False, italic=False, align=None):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run.bold = bold
    run.italic = italic
    if align is not None:
        paragraph.alignment = align
    return paragraph


def add_bullet_list(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.line_spacing = 1.5
        paragraph.paragraph_format.space_after = Pt(4)
        run = paragraph.add_run(item)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)


def add_page_footer(section):
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_page_number(paragraph)


def add_new_page(doc):
    new_section = doc.add_section(WD_SECTION.NEW_PAGE)
    add_page_footer(new_section)
    return new_section


doc = Document()
normal_style = doc.styles["Normal"]
normal_style.font.name = "Times New Roman"
normal_style.font.size = Pt(12)

section = doc.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1.25)
section.right_margin = Inches(1)
add_page_footer(section)

# Page 1: Title Page
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(18)
run = title.add_run("RESEARCH PAPER\nON\nAWANSHI ENTERPRISES INVOICE SOFTWARE")
run.font.name = "Times New Roman"
run.font.size = Pt(20)
run.bold = True

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.paragraph_format.space_after = Pt(18)
run = subtitle.add_run(
    "A Study on the Design and Development of a Browser-Based Billing and Proforma Invoice System"
)
run.font.name = "Times New Roman"
run.font.size = Pt(14)
run.italic = True

for line in [
    "",
    "Submitted in partial fulfillment of academic project documentation requirements",
    "",
    "Prepared By:",
    "Student Name: ____________________",
    "Course / Class: ____________________",
    "Institution: ____________________",
    "",
    "Guided By:",
    "Project Guide: ____________________",
    "",
    "Academic Year: 2026",
]:
    add_paragraph(doc, line, align=WD_ALIGN_PARAGRAPH.CENTER)

# Page 2: Certificate
add_new_page(doc)
add_heading(doc, "Certificate", level=1)
add_paragraph(
    doc,
    "This is to certify that the research paper titled 'AWANSHI ENTERPRISES Invoice "
    "Software' is a genuine record of work prepared for academic submission. The study "
    "focuses on the planning, analysis, design, implementation, and evaluation of a "
    "browser-based billing platform intended to support invoice generation, stock-aware "
    "item management, GST-ready calculations, and secure user-specific data handling.",
)
add_paragraph(
    doc,
    "The work reflects a practical attempt to solve real-world operational issues faced by "
    "small and medium enterprises that depend on fast invoice creation, consistent record "
    "maintenance, and simple digital workflows. The report documents the motivation of the "
    "project, the technical architecture, the tools used, and the observed outcomes after "
    "implementation and testing.",
)
add_paragraph(
    doc,
    "To the best of our knowledge, the matter included in this report has been organized "
    "for educational and project demonstration purposes and can be evaluated as a complete "
    "study of the developed system.",
)
add_paragraph(doc, "")
add_paragraph(doc, "Guide Signature: ____________________")
add_paragraph(doc, "Head of Department: ____________________")
add_paragraph(doc, "Date: ____________________")

# Page 3: Declaration
add_new_page(doc)
add_heading(doc, "Declaration", level=1)
add_paragraph(
    doc,
    "I hereby declare that this research paper on 'AWANSHI ENTERPRISES Invoice Software' "
    "is based on original project work and has been prepared for academic use. All facts, "
    "observations, explanations, and technical descriptions included in this report have "
    "been compiled carefully for the purpose of documenting a billing software solution for "
    "small business operations.",
)
add_paragraph(
    doc,
    "The project idea was shaped around the practical requirement of generating bills and "
    "proforma invoices in a fast, accurate, and user-friendly manner. The report may include "
    "references to common web technologies, online documentation, software deployment "
    "concepts, and GST-oriented billing practices, all of which have been used only to "
    "support the academic explanation of the developed system.",
)
add_paragraph(
    doc,
    "I further declare that this report has not been submitted elsewhere in the same form "
    "for any other degree, diploma, or formal certification.",
)
add_paragraph(doc, "")
add_paragraph(doc, "Student Signature: ____________________")
add_paragraph(doc, "Name: ____________________")
add_paragraph(doc, "Date: ____________________")

# Page 4: Acknowledgement
add_new_page(doc)
add_heading(doc, "Acknowledgement", level=1)
add_paragraph(
    doc,
    "The completion of this research paper and software study became possible with the "
    "guidance, support, and encouragement received throughout the project period. I express "
    "my sincere gratitude to my teacher and project guide for providing direction, valuable "
    "suggestions, and motivation during the planning and preparation of this report.",
)
add_paragraph(
    doc,
    "I also acknowledge the importance of practical business requirements that inspired the "
    "development of this billing system. The operational needs of AWANSHI ENTERPRISES helped "
    "shape the software into a realistic solution focused on invoice generation, customer "
    "record handling, stock-aware item storage, GST calculations, secure login, and backup "
    "management.",
)
add_paragraph(
    doc,
    "Finally, I thank all those who directly or indirectly supported this work through "
    "feedback, testing, review, and encouragement. Their support contributed meaningfully to "
    "the completion of both the software project and the present research paper.",
)

# Page 5: Abstract
add_new_page(doc)
add_heading(doc, "Abstract", level=1)
add_paragraph(
    doc,
    "This research paper presents the study and documentation of AWANSHI ENTERPRISES "
    "Invoice Software, a browser-based billing application created to simplify invoice and "
    "proforma invoice generation for business users. Traditional billing processes often rely "
    "on handwritten records or scattered spreadsheet-based practices that lead to repetitive "
    "manual effort, inconsistent calculations, limited stock visibility, and difficulty in "
    "retrieving past records. The proposed software addresses these challenges by offering a "
    "structured and digital workflow for invoice creation, item management, workspace-based "
    "data handling, GST calculation, preview generation, and printable outputs.",
)
add_paragraph(
    doc,
    "The system is implemented using web technologies such as HTML, CSS, and JavaScript, "
    "with Supabase used for authentication and per-user data storage. Deployment readiness is "
    "supported through a Vercel-compatible backend layer, enabling the same logic to function "
    "during local testing and hosted execution. The software includes item storage, live "
    "document preview, auto-numbering of documents, account onboarding, backup import/export, "
    "and per-user isolation of records.",
)
add_paragraph(
    doc,
    "The study concludes that the developed system improves billing accuracy, enhances "
    "record organization, and offers a practical digital foundation for small businesses. "
    "The paper also identifies future opportunities such as analytics, mobile adaptation, "
    "barcode support, and richer reporting modules.",
)

# Page 6: Introduction
add_new_page(doc)
add_heading(doc, "1. Introduction", level=1)
add_paragraph(
    doc,
    "Billing is a core operational activity in every commercial organization. Whether a "
    "business sells products or provides services, it must produce clear invoices, maintain "
    "customer records, preserve transaction history, and calculate taxes accurately. In many "
    "small businesses, these tasks are still performed with paper registers, calculators, or "
    "simple spreadsheet files. Such methods may work for a limited volume of transactions, "
    "but they become inefficient when the business grows or when records must be shared, "
    "searched, or audited.",
)
add_paragraph(
    doc,
    "AWANSHI ENTERPRISES Invoice Software has been designed as a web-based application to "
    "modernize this workflow. The software aims to reduce manual billing errors, improve data "
    "consistency, and provide users with a convenient environment for creating both bills and "
    "proforma invoices. By placing the workflow inside a browser, the system remains "
    "accessible and familiar while avoiding the complexity of heavyweight desktop software.",
)
add_paragraph(
    doc,
    "The project demonstrates how a focused business application can be developed with modern "
    "frontend and backend services while still remaining lightweight, readable, and usable for "
    "real operational needs. The software is especially relevant for micro, small, and local "
    "enterprises that want a digital billing system without adopting a large enterprise "
    "resource planning platform.",
)

# Page 7: Problem Statement and Objectives
add_new_page(doc)
add_heading(doc, "2. Problem Statement and Objectives", level=1)
add_heading(doc, "2.1 Problem Statement", level=2)
add_paragraph(
    doc,
    "Small business owners often face recurring difficulties during invoice preparation. "
    "Manual bill writing consumes time, requires repeated data entry, and increases the chance "
    "of mistakes in item totals, tax calculations, customer details, and payment notes. When "
    "documents are created in isolated files or notebooks, historical records become harder to "
    "search, compare, or back up. In addition, stock-related information and billing data are "
    "often maintained separately, which creates inconsistency between inventory and sales "
    "records.",
)
add_paragraph(
    doc,
    "The lack of a structured account-based system introduces another challenge: data from one "
    "user may be mixed with that of another, especially when multiple people access the same "
    "device or software instance. Therefore, a lightweight yet secure billing platform is "
    "needed to combine invoice generation, item management, tax-ready calculations, and "
    "user-specific record isolation in one place.",
)
add_heading(doc, "2.2 Objectives", level=2)
add_bullet_list(
    doc,
    [
        "To create a browser-based billing system for invoices and proforma invoices.",
        "To automate invoice numbering, item total calculations, GST handling, and grand totals.",
        "To maintain an item store with product rate, HSN code, and stock-related information.",
        "To support per-user account isolation through authenticated workspaces.",
        "To provide backup export and import for data safety and portability.",
        "To enable document preview, print readiness, and deployment on a public hosting platform.",
    ],
)

# Page 8: Literature Review
add_new_page(doc)
add_heading(doc, "3. Literature Review", level=1)
add_paragraph(
    doc,
    "The evolution of billing software has followed the wider shift from paper-based office "
    "administration to digital business management. Earlier approaches often centered on "
    "desktop accounting programs or spreadsheet templates. These systems improved calculation "
    "accuracy but usually required manual file management, local installation, and user "
    "discipline for maintaining structure. Over time, cloud-based tools introduced login-based "
    "access, remote storage, collaborative data availability, and easier recovery of records.",
)
add_paragraph(
    doc,
    "Research and industry practice show that small enterprises prefer systems that are easy to "
    "learn, low in cost, and fast to deploy. Heavy enterprise systems provide many advanced "
    "modules but can be difficult for smaller businesses that mainly need billing, inventory "
    "visibility, and customer document history. This has created demand for targeted business "
    "applications that solve a narrow but important operational problem effectively.",
)
add_paragraph(
    doc,
    "Modern web applications further improve feasibility by allowing software to run directly in "
    "the browser while connecting to managed backend services for identity, storage, and API "
    "processing. A combination of frontend scripting, authenticated data access, and serverless "
    "deployment is now sufficient to deliver practical line-of-business tools. The present "
    "project follows this trend by combining a simple user interface with secure online data "
    "handling and deployment-ready architecture.",
)

# Page 9: System Analysis and Requirements
add_new_page(doc)
add_heading(doc, "4. System Analysis and Requirements", level=1)
add_heading(doc, "4.1 Functional Requirements", level=2)
add_bullet_list(
    doc,
    [
        "User signup, login, email confirmation, and password reset.",
        "Creation of bills and proforma invoices with separate numbering.",
        "Item storage with name, stock, HSN code, and default rate.",
        "Invoice preview, print support, save, status change, and delete operations.",
        "Dashboard view for totals and recent document activity.",
        "Company onboarding with GST, address, phone, and bank note details.",
    ],
)
add_heading(doc, "4.2 Non-Functional Requirements", level=2)
add_bullet_list(
    doc,
    [
        "Usability through a simple browser-based interface.",
        "Data consistency across user accounts.",
        "Security through authenticated access and backend rules.",
        "Reliability through backup export/import and cloud storage.",
        "Scalability for local use and hosted deployment.",
    ],
)
add_heading(doc, "4.3 Technical Stack", level=2)
add_paragraph(
    doc,
    "The application uses HTML for structure, CSS for presentation, and JavaScript for frontend "
    "behavior. Supabase provides authentication and data storage, while a shared backend layer "
    "supports both local operation and Vercel API routing for hosted use.",
)

# Page 10: System Design
add_new_page(doc)
add_heading(doc, "5. System Design", level=1)
add_paragraph(
    doc,
    "The system is organized as a browser-first application with modular responsibilities. The "
    "user interface manages data entry, dashboard summaries, live document preview, and page "
    "navigation between billing features. The backend layer exposes shared logic so that local "
    "development and deployed execution behave consistently. This separation of concerns keeps "
    "the application easier to maintain and improves portability.",
)
add_paragraph(
    doc,
    "At the account level, each user has an isolated workspace. Core data entities include "
    "profiles, items, invoices, and document counters. Profiles store business identity and "
    "default settings. Items store product-related details and pricing. Invoices record both "
    "bill and proforma data, while counters preserve sequential numbering. This entity layout "
    "supports a practical real-world flow where the business profile and item list remain "
    "reusable across many future invoices.",
)
add_paragraph(
    doc,
    "The design also includes caching on the browser side for convenience, while treating the "
    "cloud-backed database as the source of truth for signed-in users. This hybrid approach "
    "balances usability and persistence by allowing responsive local interaction without losing "
    "the benefits of account-based storage.",
)

# Page 11: Implementation
add_new_page(doc)
add_heading(doc, "6. Implementation", level=1)
add_paragraph(
    doc,
    "Implementation of the software emphasizes direct usability and deployment readiness. The "
    "frontend presents a dashboard, item store, new document form, and preview interface. The "
    "invoice form captures client details, invoice number, dates, line items, GST rate, "
    "discount, notes, and status. Totals are calculated dynamically so the user can immediately "
    "verify the financial values before saving the document.",
)
add_paragraph(
    doc,
    "The backend integration supports user registration and authenticated data access. Email "
    "confirmation and password reset flows are included, which improves readiness for real user "
    "accounts. Per-user storage ensures that each account accesses only its own invoices, items, "
    "profile information, and counters. Shared backend logic allows the same project to work in "
    "local testing and hosted deployment environments.",
)
add_paragraph(
    doc,
    "The project further includes backup export and import in JSON form. This feature is "
    "important because it gives users control over their records beyond the hosted database. "
    "Operational continuity improves when business data can be archived, transferred, or "
    "restored without deep technical intervention.",
)

# Page 12: Testing and Results
add_new_page(doc)
add_heading(doc, "7. Testing and Results", level=1)
add_paragraph(
    doc,
    "Testing of the system focused on practical workflows rather than abstract edge cases alone. "
    "The major scenarios included account signup and login, item creation, invoice generation, "
    "document preview, invoice status changes, deletion, data backup, and retrieval of prior "
    "records. Special attention was given to verifying that calculations update correctly when "
    "quantity, rate, GST, or discount values change.",
)
add_paragraph(
    doc,
    "Results indicate that the application successfully reduces repeated manual effort during "
    "billing. The dashboard offers quick visibility into recent activity, while the item store "
    "supports faster document preparation by reusing stored product details. Document preview "
    "improves confidence before printing, and authenticated storage protects data separation "
    "between accounts. Together, these behaviors confirm that the developed solution meets the "
    "main operational goals identified in the problem statement.",
)
table = doc.add_table(rows=1, cols=3)
table.style = "Table Grid"
hdr = table.rows[0].cells
hdr[0].text = "Test Area"
hdr[1].text = "Expected Result"
hdr[2].text = "Observed Result"
rows = [
    ("Login and Signup", "Secure account access", "Working successfully"),
    ("Item Store", "Save and reuse product data", "Working successfully"),
    ("Invoice Calculation", "Correct total and GST values", "Working successfully"),
    ("Preview and Print", "Readable output before saving", "Working successfully"),
    ("Backup/Import", "Data portability", "Working successfully"),
]
for test_area, expected, observed in rows:
    row = table.add_row().cells
    row[0].text = test_area
    row[1].text = expected
    row[2].text = observed

# Page 13: Conclusion and Future Scope
add_new_page(doc)
add_heading(doc, "8. Conclusion and Future Scope", level=1)
add_heading(doc, "8.1 Conclusion", level=2)
add_paragraph(
    doc,
    "The research and development of AWANSHI ENTERPRISES Invoice Software demonstrate that a "
    "focused billing application can significantly improve the efficiency of day-to-day business "
    "operations. By integrating invoice preparation, item storage, account-based data handling, "
    "GST support, and deployment-ready backend logic, the project provides a practical and "
    "usable digital solution for small enterprises. The application reduces manual repetition, "
    "improves calculation accuracy, and creates a better record management workflow.",
)
add_heading(doc, "8.2 Future Scope", level=2)
add_bullet_list(
    doc,
    [
        "Advanced reporting and sales analytics dashboards.",
        "Barcode and QR code support for item handling and payment references.",
        "Mobile-responsive enhancements and packaged mobile application support.",
        "Customer statement generation and recurring invoice scheduling.",
        "Role-based access for multiple staff members under one business account.",
    ],
)

# Page 14: References
add_new_page(doc)
add_heading(doc, "References", level=1)
references = [
    "MDN Web Docs. HTML, CSS, and JavaScript reference materials for frontend development.",
    "Supabase Documentation. Authentication, database, and storage workflow references.",
    "Vercel Documentation. Deployment patterns and serverless API route hosting guidance.",
    "Business billing and invoice format guidelines commonly used in GST-ready commercial workflows.",
    "Project repository materials: README.md, app.js, index.html, styles.css, backend files, and schema notes.",
]
for index, item in enumerate(references, start=1):
    add_paragraph(doc, f"{index}. {item}")

doc.save(OUTPUT_PATH)
print(OUTPUT_PATH)
